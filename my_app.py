from docx import Document
from docx.shared import Inches
import pyttsx3

def speak(text):
    pyttsx3.speak(text)

document = Document() 

# profile picture 
document.add_picture(
    'me.jpeg', 
    width=Inches(2.0)
)

# name phone number and details 
speak('Quel est votre nom')
name = input('What is your name ? ')
speak('quel est votre numéro de téléphone')
phone_number = input('what is your phone number ? ')
speak('quel est votre email ')
email = input('What is your email? ')

document.add_paragraph(
    name + ' | ' + phone_number + ' | ' + email
)

#  about me
document.add_heading('About me')
speak('que dois-je savoir sur vous')
document.add_paragraph(
    input('Tell me about yourself ? ')
    )

# work experience
document.add_heading('Work experiences')
p = document.add_paragraph()

speak('Entrez nom entreprise')
company = input('Enter company ')
speak('date de début de contrat')
from_date = input('From Date ')
speak('date de fin de contrat')
to_date = input('To Date ')

p.add_run(company + ' ').bold = True 
p.add_run(from_date + '-' + to_date + '\n').italic = True

speak('décrivez votre expérience au sein de cette entreprise')
experience_details = input(
    'Describe your experience at ' + company + ' -docx')
p.add_run(experience_details)

# more experiences
while True:
    speak('Avez-vous une autre expérience')
    has_more_experiences = input(
        'Do you have more experiences ? Yes or No ')
    if has_more_experiences.lower() == 'yes': 
        p = document.add_paragraph()

        speak('Entrez nom entreprise')
        company = input('Enter company ')
        speak('date de début de contrat')
        from_date = input('From Date ')
        speak('date de fin de contrat')    
        to_date = input('To Date ')

        p.add_run(company + ' ').bold = True 
        p.add_run(from_date + '-' + to_date + '\n').italic = True
        
        speak('décrivez votre expérience au sein de cette entreprise')
        experience_details = input(
            'Describe your experience at ' + company + ' ')
        p.add_run(experience_details)
    else:
        break

# Skills
document.add_heading('Skills')
speak('Entrez une compétence')
skill = input('Enter skill ')
p = document.add_paragraph(skill)
p.style = 'List Bullet'

while True:
    speak('avez-vous une autre compétence?')
    has_more_skills = input('Do you have more skills ? Yes or No ')
    if has_more_skills.lower() == 'yes':
        speak('Entrez une compétence')
        skill = input('Enter skill ')
        p = document.add_paragraph(skill)
        p.style = 'List Bullet'
    else:
        speak('Merci, et à bientôt')
        break

            


    # footer
    section = document.sections[0]
    footer =  section.footer 
    p = footer.paragraphs[0]
    p.text = "CV genrated using Amigoscode and Intuit Quickbooks course project"

    
document.save('cv.docx')